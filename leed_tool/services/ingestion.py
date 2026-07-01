"""Safe, in-memory extraction for uploaded project documents and drawings."""

from __future__ import annotations

import csv
import hashlib
import mimetypes
import re
from dataclasses import dataclass, field
from io import BytesIO, StringIO
from pathlib import Path, PurePosixPath
from zipfile import BadZipFile, ZipFile

from docx import Document
from openpyxl import load_workbook
from PIL import Image
from pypdf import PasswordType, PdfReader
from pypdf.errors import DependencyError, PyPdfError

try:  # Optional high-fidelity drawing parsers installed in production.
    import fitz  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - local lightweight test environment
    fitz = None

try:
    import ezdxf  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - local lightweight test environment
    ezdxf = None

MAX_FILE_BYTES = 60 * 1024 * 1024
MAX_ARCHIVE_BYTES = 500 * 1024 * 1024
MAX_ARCHIVE_FILES = 250
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 1024 * 1024 * 1024
MAX_ARCHIVE_COMPRESSION_RATIO = 200
MAX_TEXT_CHARS = 250_000
MAX_PDF_PAGES = 300
FAST_PDF_PAGES = 10
FAST_PDF_BYTES = 2 * 1024 * 1024
MAX_FAST_PARSED_PDFS = 50
MAX_SHEET_ROWS = 5_000
MAX_SHEET_COLUMNS = 80

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".xlsx", ".xlsm", ".csv", ".txt", ".md",
    ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".dxf", ".ifc", ".dwg",
}


@dataclass(frozen=True, slots=True)
class DocumentRecord:
    name: str
    extension: str
    content_type: str
    size_bytes: int
    sha256: str
    text: str
    page_count: int = 0
    kind: str = "Document"
    extraction_status: str = "Extracted"
    warnings: tuple[str, ...] = field(default_factory=tuple)


def _safe_name(name: str) -> str:
    return Path(name.replace("\x00", "")).name[:180] or "unnamed-file"


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _truncate(text: str, warnings: list[str]) -> str:
    clean = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(clean) > MAX_TEXT_CHARS:
        warnings.append(f"Extracted text was limited to {MAX_TEXT_CHARS:,} characters for online analysis.")
        return clean[:MAX_TEXT_CHARS]
    return clean


def _extract_pdf(
    data: bytes, warnings: list[str], page_limit: int = MAX_PDF_PAGES
) -> tuple[str, int]:
    try:
        reader = PdfReader(BytesIO(data), strict=False)
    except DependencyError as exc:
        raise ValueError(
            "Encrypted PDF support is unavailable on the server. "
            "Install pypdf[crypto] or upload an unlocked PDF."
        ) from exc
    except PyPdfError as exc:
        raise ValueError(f"PDF structure could not be read: {exc}") from exc

    if reader.is_encrypted:
        try:
            password_type = reader.decrypt("")
        except DependencyError as exc:
            raise ValueError(
                "Encrypted PDF support is unavailable on the server. "
                "Install pypdf[crypto] or upload an unlocked PDF."
            ) from exc
        except PyPdfError as exc:
            raise ValueError("Password-protected PDF cannot be analyzed; upload an unlocked copy.") from exc
        if password_type == PasswordType.NOT_DECRYPTED:
            raise ValueError("Password-protected PDF cannot be analyzed; upload an unlocked copy.")
    page_count = len(reader.pages)
    if page_count > page_limit:
        warnings.append(f"Only the first {page_limit} of {page_count} PDF pages were analyzed.")
    pages: list[str] = []
    for index, page in enumerate(reader.pages[:page_limit], start=1):
        try:
            extracted = page.extract_text() or ""
        except Exception:
            extracted = ""
            warnings.append(f"Page {index} text extraction failed.")
        if extracted.strip():
            pages.append(f"[Page {index}]\n{extracted}")
    if fitz is not None:
        try:
            drawing = fitz.open(stream=data, filetype="pdf")
            vector_pages: list[str] = []
            image_count = 0
            for page_index in range(min(len(drawing), page_limit)):
                page = drawing.load_page(page_index)
                index = page_index + 1
                blocks = page.get_text("blocks")
                block_text = "\n".join(str(block[4]).strip() for block in blocks if str(block[4]).strip())
                image_count += len(page.get_images(full=True))
                if block_text:
                    vector_pages.append(f"[Drawing page {index} text blocks]\n{block_text}")
            drawing.close()
            vector_text = "\n".join(vector_pages)
            pypdf_text = "\n".join(pages)
            if len(vector_text) > len(pypdf_text) * 1.1:
                pages = vector_pages
                warnings.append("Vector drawing text blocks were extracted with the enhanced PDF parser.")
            if image_count:
                warnings.append(f"PDF contains {image_count} embedded image(s); raster-only annotations may require OCR or visual review.")
        except Exception as exc:  # pragma: no cover - optional parser differences
            warnings.append(f"Enhanced drawing parser fallback used: {type(exc).__name__}.")
    if not pages:
        warnings.append("No searchable PDF text was found; this file likely needs OCR or visual review.")
    return "\n".join(pages), page_count


def _extract_docx(data: bytes) -> str:
    document = Document(BytesIO(data))
    blocks = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(blocks)


def _extract_xlsx(data: bytes, warnings: list[str]) -> str:
    workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
    blocks: list[str] = []
    try:
        for sheet in workbook.worksheets:
            blocks.append(f"[Worksheet: {sheet.title}]")
            for row_index, row in enumerate(
                sheet.iter_rows(max_row=MAX_SHEET_ROWS, max_col=MAX_SHEET_COLUMNS, values_only=True),
                start=1,
            ):
                values = [str(value) for value in row if value is not None]
                if values:
                    blocks.append(" | ".join(values))
                if row_index >= MAX_SHEET_ROWS:
                    warnings.append(f"Worksheet '{sheet.title}' was limited to {MAX_SHEET_ROWS:,} rows.")
                    break
    finally:
        workbook.close()
    return "\n".join(blocks)


def _extract_csv(data: bytes) -> str:
    decoded = _decode_text(data)
    rows = csv.reader(StringIO(decoded))
    output: list[str] = []
    for index, row in enumerate(rows, start=1):
        if index > MAX_SHEET_ROWS:
            break
        output.append(" | ".join(row[:MAX_SHEET_COLUMNS]))
    return "\n".join(output)


def _extract_dxf(data: bytes, warnings: list[str]) -> str:
    decoded = _decode_text(data)
    if ezdxf is None:
        warnings.append("DXF text was read without the optional structured CAD parser.")
        return decoded
    try:
        drawing = ezdxf.read(StringIO(decoded))
        modelspace = drawing.modelspace()
        layers: set[str] = set()
        blocks: set[str] = set()
        annotations: list[str] = []
        entity_counts: dict[str, int] = {}
        for entity in modelspace:
            entity_type = entity.dxftype()
            entity_counts[entity_type] = entity_counts.get(entity_type, 0) + 1
            if hasattr(entity.dxf, "layer"):
                layers.add(str(entity.dxf.layer))
            if entity_type == "INSERT":
                blocks.add(str(entity.dxf.name))
            elif entity_type == "TEXT":
                annotations.append(str(entity.dxf.text))
            elif entity_type == "MTEXT":
                annotations.append(entity.plain_text())
        summary = [
            "[Structured DXF summary]",
            "Layers: " + ", ".join(sorted(layers)[:200]),
            "Blocks: " + ", ".join(sorted(blocks)[:200]),
            "Entity counts: " + ", ".join(f"{key}={value}" for key, value in sorted(entity_counts.items())),
            "Annotations:",
            *annotations[:5000],
        ]
        return "\n".join(summary)
    except Exception as exc:
        warnings.append(f"Structured DXF parsing failed ({type(exc).__name__}); raw text fallback used.")
        return decoded


def _extract_ifc(data: bytes, warnings: list[str]) -> str:
    decoded = _decode_text(data)
    entity_types = re.findall(r"=\s*(IFC[A-Z0-9_]+)\s*\(", decoded, re.IGNORECASE)
    counts: dict[str, int] = {}
    for entity_type in entity_types:
        key = entity_type.upper()
        counts[key] = counts.get(key, 0) + 1
    names = re.findall(
        r"IFC(?:SPACE|ZONE|EQUIPMENTELEMENT|FURNISHINGELEMENT|BUILDINGELEMENTPROXY)\s*\([^;]{0,500}?'([^']{2,100})'",
        decoded,
        re.IGNORECASE,
    )
    warnings.append("IFC STEP entities and labels were indexed; geometric compliance still requires a BIM authoring/model-checking workflow.")
    summary = [
        "[Structured IFC index]",
        "Entity counts: " + ", ".join(f"{key}={value}" for key, value in sorted(counts.items())[:250]),
        "Detected space/equipment labels: " + ", ".join(dict.fromkeys(names[:1000])),
        "[Raw IFC text]",
        decoded,
    ]
    return "\n".join(summary)


def _inspect_image(data: bytes, name: str, warnings: list[str]) -> str:
    with Image.open(BytesIO(data)) as image:
        width, height = image.size
        image_format = image.format or "raster"
        image.verify()
    warnings.append("Raster drawing validated, but visual content requires OCR/vision review; use searchable PDF for full automated evidence extraction.")
    return f"Raster drawing file {name}. Format {image_format}. Image dimensions {width} x {height} pixels."


def extract_document(
    name: str,
    data: bytes,
    content_type: str = "",
    deep_scan: bool = False,
    inventory_only: bool = False,
) -> DocumentRecord:
    """Validate and extract one upload without writing it to disk."""

    safe_name = _safe_name(name)
    extension = Path(safe_name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {extension or 'no extension'}")
    if not data:
        raise ValueError("The uploaded file is empty.")
    if len(data) > MAX_FILE_BYTES:
        raise ValueError(f"File exceeds the {MAX_FILE_BYTES // (1024 * 1024)} MB online limit.")

    warnings: list[str] = []
    page_count = 0
    kind = "Document"
    status = "Extracted"
    try:
        if extension == ".pdf":
            kind = "Drawing / PDF"
            if inventory_only or (not deep_scan and len(data) > FAST_PDF_BYTES):
                text = f"Large drawing or scanned PDF inventoried for visual review: {safe_name}."
                status = "Visual review required"
                warnings.append(
                    f"Rapid mode skipped full text extraction for this {len(data) / (1024 * 1024):.1f} MB PDF. "
                    "Use Deep drawing scan for page-level extraction."
                )
            else:
                text, page_count = _extract_pdf(
                    data, warnings, MAX_PDF_PAGES if deep_scan else FAST_PDF_PAGES
                )
        elif extension == ".docx":
            text = _extract_docx(data)
            kind = "Specification / Narrative"
        elif extension in {".xlsx", ".xlsm"}:
            text = _extract_xlsx(data, warnings)
            kind = "Calculation / Schedule"
            if extension == ".xlsm":
                warnings.append("VBA macros were not executed or analyzed; only cached worksheet values were read.")
        elif extension == ".csv":
            text = _extract_csv(data)
            kind = "Calculation / Schedule"
        elif extension in {".txt", ".md"}:
            text = _decode_text(data)
            kind = "Narrative"
        elif extension == ".dxf":
            text = _extract_dxf(data, warnings)
            kind = "BIM / CAD Export"
        elif extension == ".ifc":
            text = _extract_ifc(data, warnings)
            kind = "BIM / CAD Export"
        elif extension in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
            text = _inspect_image(data, safe_name, warnings)
            kind = "Raster Drawing"
            status = "Visual review required"
        else:  # DWG is intentionally accepted into the manifest but not decoded.
            text = f"Binary CAD drawing {safe_name}."
            kind = "CAD Drawing"
            status = "Conversion required"
            warnings.append("Native DWG is inventoried but not parsed; export to searchable PDF, IFC, or DXF for automated review.")
    except (ValueError, OSError, KeyError, EOFError, DependencyError, PyPdfError) as exc:
        raise ValueError(f"Could not safely parse {safe_name}: {exc}") from exc

    text = _truncate(text, warnings)
    return DocumentRecord(
        name=safe_name,
        extension=extension,
        content_type=content_type or "application/octet-stream",
        size_bytes=len(data),
        sha256=hashlib.sha256(data).hexdigest(),
        text=text,
        page_count=page_count,
        kind=kind,
        extraction_status=status,
        warnings=tuple(dict.fromkeys(warnings)),
    )


def _archive_member_path(filename: str) -> PurePosixPath | None:
    """Return a normalized safe member path without extracting it to disk."""

    normalized = filename.replace("\\", "/").replace("\x00", "")
    path = PurePosixPath(normalized)
    if not normalized or path.is_absolute() or ".." in path.parts:
        return None
    return path


def _archive_member_label(archive_name: str, path: PurePosixPath) -> str:
    archive = _safe_name(archive_name)
    member = " › ".join(part for part in path.parts if part not in {"", "."})
    label = f"{archive} :: {member}"
    if len(label) <= 180:
        return label
    member_tail = member if len(member) <= 105 else f"{member[:35]}…{member[-69:]}"
    archive_head = archive[: max(30, 174 - len(member_tail))]
    return f"{archive_head}… :: {member_tail}"


def _is_priority_archive_member(path: PurePosixPath) -> bool:
    lower = path.as_posix().lower()
    return any(
        token in lower
        for token in (
            "review report", "review_report", "scorecard", "clarification", "appeal",
            "narrative", "calculator", "summary report", "simulation summary",
            "commissioning report", "owner statement", "commitment", "leed form",
        )
    )


def _extract_zip_archive(
    name: str, data: bytes, depth: int = 0, deep_scan: bool = False
) -> tuple[list[DocumentRecord], list[str]]:
    """Safely inspect and parse supported files from one ZIP archive in memory."""

    archive_name = _safe_name(name)
    if not data:
        return [], [f"Could not safely parse {archive_name}: the uploaded archive is empty."]
    if len(data) > MAX_ARCHIVE_BYTES:
        return [], [
            f"Could not safely parse {archive_name}: ZIP exceeds the "
            f"{MAX_ARCHIVE_BYTES // (1024 * 1024)} MB compressed upload limit."
        ]

    documents: list[DocumentRecord] = []
    errors: list[str] = []
    skipped: dict[str, list[str]] = {
        "unsafe paths": [],
        "unsupported types": [],
        "nested ZIP archives": [],
        "encrypted members": [],
        "oversized members": [],
        "suspicious compression ratios": [],
    }
    try:
        with ZipFile(BytesIO(data)) as archive:
            members = [info for info in archive.infolist() if not info.is_dir()]
            members.sort(
                key=lambda info: (
                    0 if _is_priority_archive_member(PurePosixPath(info.filename.replace("\\", "/"))) else 1,
                    0 if PurePosixPath(info.filename).suffix.lower() in {".xlsm", ".xlsx", ".docx", ".csv", ".txt", ".md"} else 1,
                    info.file_size,
                )
            )
            total_uncompressed = sum(info.file_size for info in members)
            if total_uncompressed > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
                return [], [
                    f"Could not safely parse {archive_name}: ZIP expands to "
                    f"{total_uncompressed / (1024 * 1024):.1f} MB; the safe limit is "
                    f"{MAX_ARCHIVE_UNCOMPRESSED_BYTES // (1024 * 1024)} MB."
                ]
            if len(members) > MAX_ARCHIVE_FILES:
                errors.append(
                    f"{archive_name}: only the first {MAX_ARCHIVE_FILES} of "
                    f"{len(members)} archive members were inspected."
                )

            rapid_parsed_pdfs = 0
            for info in members[:MAX_ARCHIVE_FILES]:
                path = _archive_member_path(info.filename)
                display_name = info.filename[:120] or "unnamed member"
                if path is None:
                    skipped["unsafe paths"].append(display_name)
                    continue
                if info.flag_bits & 0x1:
                    skipped["encrypted members"].append(display_name)
                    continue
                extension = path.suffix.lower()
                if info.file_size > MAX_FILE_BYTES:
                    skipped["oversized members"].append(display_name)
                    continue
                ratio = info.file_size / max(info.compress_size, 1)
                if info.file_size > 1024 * 1024 and ratio > MAX_ARCHIVE_COMPRESSION_RATIO:
                    skipped["suspicious compression ratios"].append(display_name)
                    continue

                member_name = _archive_member_label(archive_name, path)
                if extension == ".zip":
                    if depth >= 1:
                        skipped["nested ZIP archives"].append(display_name)
                        continue
                    try:
                        nested_documents, nested_errors = _extract_zip_archive(
                            member_name, archive.read(info), depth + 1, deep_scan
                        )
                        documents.extend(nested_documents)
                        errors.extend(nested_errors)
                    except (RuntimeError, ValueError, OSError, KeyError, EOFError) as exc:
                        errors.append(f"{member_name}: {exc}")
                    continue
                if extension not in SUPPORTED_EXTENSIONS:
                    skipped["unsupported types"].append(display_name)
                    continue
                try:
                    member_data = archive.read(info)
                    content_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
                    priority_member = _is_priority_archive_member(path)
                    priority_parse = priority_member and info.file_size <= 8 * 1024 * 1024
                    member_deep_scan = deep_scan or priority_parse
                    inventory_only = False
                    if extension == ".pdf" and not deep_scan:
                        would_parse = priority_parse or info.file_size <= FAST_PDF_BYTES
                        if would_parse and rapid_parsed_pdfs < MAX_FAST_PARSED_PDFS:
                            rapid_parsed_pdfs += 1
                        else:
                            inventory_only = True
                    documents.append(
                        extract_document(
                            member_name, member_data, content_type, member_deep_scan, inventory_only
                        )
                    )
                except (RuntimeError, ValueError, OSError, KeyError, EOFError) as exc:
                    errors.append(f"{member_name}: {exc}")
    except (BadZipFile, OSError, EOFError) as exc:
        return [], [f"Could not safely parse {archive_name}: invalid or damaged ZIP archive ({exc})."]

    for reason, names in skipped.items():
        if names:
            examples = ", ".join(names[:3])
            suffix = "" if len(names) <= 3 else f" and {len(names) - 3} more"
            errors.append(
                f"{archive_name}: skipped {len(names)} member(s) due to {reason}: {examples}{suffix}."
            )
    if not documents and not errors:
        errors.append(f"{archive_name}: the ZIP archive contains no supported project files.")
    return documents, errors


def extract_many(
    files: list[tuple[str, bytes, str]], deep_scan: bool = False
) -> tuple[list[DocumentRecord], list[str]]:
    """Extract uploads and ZIP packages while isolating failures to individual files."""

    documents: list[DocumentRecord] = []
    errors: list[str] = []
    for name, data, content_type in files:
        if Path(_safe_name(name)).suffix.lower() == ".zip":
            archive_documents, archive_errors = _extract_zip_archive(name, data, deep_scan=deep_scan)
            documents.extend(archive_documents)
            errors.extend(archive_errors)
            continue
        try:
            documents.append(extract_document(name, data, content_type, deep_scan))
        except ValueError as exc:
            errors.append(str(exc))
    return documents, errors
